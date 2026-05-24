# -*- coding: cp1254 -*-
from pathlib import Path
from docx import Document


SRC = next(p for p in Path(".").glob("*BT_revize.docx") if not p.name.startswith("~$"))
OUT = Path("BT_revize_analitik_detayli_TURKCE.docx")


section = [
    ("h2", "3.4. Patoloji Bazlı Sentez"),
    ("p", "Bu bölümde patoloji grupları yalnızca çalışma sayılarıyla değil, çalışmaların hangi klinik soruya yanıt aradığı, hangi bilgisayarlı görme görevlerine yoğunlaştığı, hangi model ailelerini kullandığı ve kanıtın klinik uygulamaya ne kadar yaklaştığı açısından sentezlenmiştir. Her alt başlıkta tüm çalışmalar tek tek sıralanmamış; bunun yerine çalışmalar tematik alt kümeler halinde yorumlanmış ve bu alt kümeleri temsil eden çalışmalar ayrıntılandırılmıştır."),
    ("h3", "3.4.1. Ürolitiyazis/Nefrolitiyazis"),
    ("p", "Ürolitiyazis/nefrolitiyazis 95 çalışma ile en geniş kanıt alanıdır. Çalışmalar taş varlığını saptama, sınıflandırma, segmentasyon/lokalizasyon ve klinik risk ya da tedavi sonucuna yönelik öngörü alt kümelerinde toplanmaktadır. Saptama ve sınıflandırmanın baskın olması, bu alandaki klinik ihtiyacın hızlı tanı ve iş akışı triyajı ile ilişkili olduğunu göstermektedir."),
    ("p", "Cha ve arkadaşları şüpheli ürolitiyaziste açıklanabilir makine öğrenmesi ile bireyselleştirilmiş tanısal akıl yürütmeyi ele almıştır [4]. Chen HW ve arkadaşları klinik ve idrar parametreleriyle klinik olarak anlamlı nefrolitiyazis taramasını çok merkezli bir makine öğrenmesi modeliyle değerlendirmiştir [5]. Dimlo, Katti, Lilhore, Kumari ve Deshmukh gibi çalışmalar BT’de taş saptama ve sınıflandırma için hibrit derin öğrenme, CNN ve transformer tabanlı yaklaşımların yoğunlaştığını göstermektedir [7,11,13,15,17]."),
    ("p", "Segmentasyon ve hacimsel analiz ekseninde Gorli ve Dash, Vision Transformer ve 3B U-Net bileşenleriyle taş, kist ve tümör segmentasyonunu birlikte ele almıştır [9]. Mahmud, Perez Martinez, Velusamy, Öksüz, Choi ve Kaviani gibi çalışmalar taş segmentasyonu, lokalizasyon ve bilgisayar destekli tanı görevlerinin nicel karar desteğine doğru genişlediğini göstermektedir [14,19,22,32,46,61]."),
    ("p", "Öngörü çalışmaları daha heterojendir ancak klinik değer açısından önemlidir. Ahmad, Shafiq, Gupta, Huang, Zeng, Kim ve Yang gibi çalışmalar taş riski, tedavi sonucu, kompozisyon veya klinik karar desteği için makine öğrenmesi/radyomik modelleri kullanmıştır [33,34,48,53,105,131,159]. Buna karşın 95 çalışmanın yalnızca 14’ünde harici doğrulama kodlanmıştır; bu nedenle bu alan nicel olarak güçlü, fakat genellenebilirlik açısından halen sınırlıdır."),
    ("h3", "3.4.2. Abdominal Aort Anevrizması/Diseksiyonu"),
    ("p", "Abdominal aort anevrizması/diseksiyonu 79 çalışma ile ikinci en büyük kanıt alanıdır. Çalışmalar segmentasyon, otomatik çap/hacim ölçümü, diseksiyon veya akut aort sendromu saptama, endoleak değerlendirmesi, büyüme öngörüsü ve triyaj alt kümelerinde yoğunlaşmıştır. Bu nedenle aort literatürü, YZ’nin yalnızca tanı değil kantitatif izlem ve risk öngörüsü aracı olarak kullanıldığını göstermektedir."),
    ("p", "Arampatzis ve arkadaşları abdominal aort anevrizması segmentasyonu için denetimsiz ve derin öğrenme yöntemlerini karşılaştırmıştır [2]. Pouncey ve arkadaşları otomatik hacim segmentasyonu ile anevrizma büyümesini değerlendirmiş, Roby ve arkadaşları segmentasyonu hasta-özel duvar stresi kestirimiyle birleştirmiştir [18,20]. Cai, Zhang, Norajitra, Zhuang, Yu ve Wobben gibi çalışmalar aort lümeni, yalancı lümen, tromboz ve diseksiyon segmentasyonunda U-Net, nnU-Net, SAM ve CNN tabanlı yaklaşımların öne çıktığını göstermektedir [3,29,76,107,230,232]."),
    ("p", "Triyaj ve acil tanı ekseninde Hata ve arkadaşları kontrastsız BT’de aort diseksiyonu saptama algoritması geliştirmiş, Guo ve arkadaşları kontrastsız BT tabanlı radyomik imza ile torasik diseksiyon taramasını çok merkezli olarak değerlendirmiştir [225,226]. Tang, Wu, Chang, Laletin, Liu ve Hu gibi çalışmalar akut aort sendromu ve diseksiyon saptamasının acil iş akışı için yüksek potansiyel taşıdığını göstermektedir [21,25,43,52,133,137]."),
    ("p", "Bu grupta 20/79 harici doğrulama kodlanmıştır; bu oran diğer patolojilere göre daha güçlüdür. Yine de çalışmaların önemli kısmı teknik segmentasyon metriklerine odaklandığından, raporlama süresi, triyaj kararı, cerrahi/endovasküler planlama ve hasta sonuçları üzerindeki etki daha fazla prospektif çalışma gerektirir."),
    ("h3", "3.4.3. Akut Pankreatit"),
    ("p", "Akut pankreatit 33 çalışma ile orta büyüklükte, fakat klinik olarak anlamlı bir kanıt alanıdır. Bu grupta çalışmalar tanıdan çok şiddet, prognoz, rekürrens, nekroz, koleksiyon ve komplikasyon öngörüsüne yönelmiştir. Bu özellik, pankreatit alanını erken risk katmanlandırması için önemli hale getirmektedir."),
    ("p", "Radyomik ve klasik makine öğrenmesi çalışmaları belirgindir. Li ve arkadaşları belirsiz pankreatik kanal dilatasyonu gibi klinik tabloların değerlendirilmesine katkıda bulunmuş, Nalliah ve arkadaşları BT tabanlı radyomik modeli pankreatit sınıflandırmasında kullanmıştır [12,16]. Chen H, Bette, Yin, Zhang ve Zhou gibi çalışmalar radyomik, klinik değişken ve derin öğrenme özelliklerini birleştirerek prognoz, şiddet veya komplikasyon öngörüsüne yönelmiştir [30,45,112,160,162]."),
    ("p", "Derin öğrenme ekseninde Güneri ve arkadaşları Vision Transformer ile akut pankreatit-normal pankreas ayrımını değerlendirmiştir [10]. Xu ve arkadaşları abdominal BT’den akut pankreatit şiddetini öngören model bildirmiş, Wan ve arkadaşları rekürrens öngörüsü için multimodal derin öğrenme modeli sunmuştur [24,26]. Gupta, Kulkarni, Tartari, Wang ve Hong gibi çalışmalar ise sıvı koleksiyonu, segmentasyon ve pankreatik/peripankreatik bulguların otomatik analizine katkıda bulunmuştur [49,50,97,102,125]."),
    ("p", "Pankreatit alanında 9/33 çalışmada harici doğrulama vardır. Ancak hedef değişkenlerin tanı, şiddet, nekroz, rekürrens ve prognoz gibi farklı uçlara dağılması, çalışmalar arası karşılaştırmayı sınırlamaktadır."),
    ("h3", "3.4.4. Akut Apandisit"),
    ("p", "Akut apandisit grubunda 14 çalışma vardır. Çalışmalar apendiks saptama, akut apandisit sınıflandırması, lokalizasyon, komplike/nonkomplike ayrım ve karar destek görevlerine odaklanmıştır. Buna karşın harici doğrulama kodlanan çalışma bulunmaması en önemli sınırlılıktır."),
    ("p", "Huang ve arkadaşları 3B BT’de apandisit sınıflandırması için hiyerarşik kesit dikkat ağı geliştirmiş, Kim ve arkadaşları otomatik 3B CNN modeliyle apandisit değerlendirmesini ele almıştır [41,62]. Takaishi ve Done gibi çalışmalar lokalizasyon ve saptama görevlerine, Tsai ve Hariri gibi çalışmalar ise abdominal hastalık lokalizasyonu ve dual-path CNN tabanlı tanı yaklaşımlarına katkıda bulunmuştur [55,94,121,123]."),
    ("p", "Klasik makine öğrenmesi ve radyomik çalışmaları daha çok karar destek ve komplike apandisit ayrımına yöneliktir. An ve arkadaşları otomatik makine öğrenmesi ve özellik mühendisliğini değerlendirmiş, Zhao ve arkadaşları BT radyomik özellikleri ile klinik bilgiyi birleştirerek basit/basit olmayan apandisit ayrımına odaklanmıştır [110,164]. Baştüğ ve Ng gibi çalışmalar U-Net ve segmentasyon temelli yaklaşımların apendiks saptamada kullanılabileceğini göstermiştir [111,193]."),
    ("h3", "3.4.5. Akut Kolesistit"),
    ("p", "Akut kolesistit grubu yalnızca 5 çalışma ile sınırlı ve heterojen bir alandır. Çalışmalar akut kolesistit tanısı, suppuratif kolesistit öngörüsü, zor laparoskopik kolesistektomi öngörüsü, safra kesesi segmentasyonu/görüntü reçetelendirme ve ayırıcı tanı uçlarına dağılmıştır."),
    ("p", "Chen BQ ve arkadaşları derin öğrenmeyle akut kolesistit tanısı ve suppuratif kolesistit öngörüsünü ele almıştır [44]. Sun ve arkadaşları BT tabanlı radyomik-klinik modelle zor laparoskopik kolesistektomi öngörüsünü değerlendirmiş, Yang ve arkadaşları safra kesesi BT görüntü reçetelendirmesini otomatikleştiren derin öğrenme yaklaşımı geliştirmiştir [92,103]. Zhang ve Fujita gibi çalışmalar ise ksantogranülomatöz kolesistit ve safra kesesi tümörü ayrımı gibi ayırıcı tanı senaryolarına katkıda bulunmaktadır [163,206]."),
    ("p", "Bu grupta 1/5 çalışmada harici doğrulama vardır. Kanıt hacmi düşük ve klinik uçlar heterojen olduğundan, akut kolesistit özelinde güçlü genelleme yapmak için henüz erkendir."),
    ("h3", "3.4.6. Akut Divertikülit"),
    ("p", "Akut divertikülit 4 çalışma ile en az temsil edilen patolojilerden biridir. Çalışmalar doğrudan akut divertikülit tanısından çok sigmoid kolon lokalizasyonu/segmentasyonu, divertikülit-kolon kanseri ayrımı ve cerrahi süre öngörüsü gibi dolaylı ya da ayırıcı tanı odaklı görevlere yönelmiştir."),
    ("p", "Rahman ve arkadaşları 3B CNN ile akut divertikülit bağlamında sigmoid kolon lokalizasyonunu ele almış, M. A. Rahman ve arkadaşları 3B BT’de sigmoid kolon segmentasyonunu değerlendirmiştir [80,184]. Lippenberger ve arkadaşları görüntü tabanlı Random Forest sınıflandırıcı ile cerrahi süre öngörüsüne yönelmiş, Ziegelmayer ve arkadaşları kolon karsinomu-divertikülit ayrımını derin öğrenme ile değerlendirmiştir [136,198]."),
    ("p", "Bu grupta harici doğrulama kodlanmamıştır. Çalışma sayısının azlığı nedeniyle divertikülit alanında performans veya klinik uygulanabilirlik hakkında güçlü sonuç çıkarılamaz."),
    ("h3", "3.4.7. Birden Fazla Hedef Patoloji İçeren Çalışmalar"),
    ("p", "Birden fazla hedef patoloji içeren 4 çalışma, tek hastalık odaklı modellerden çok geniş abdominal karar destek veya çoklu patoloji saptama yaklaşımına geçişi temsil etmektedir. Bu yaklaşım klinik gerçekliğe daha yakındır; çünkü acil serviste BT yorumlama süreci çoğu zaman birden fazla olası patolojiyi dışlama veya önceliklendirme problemidir."),
    ("p", "Li ve arkadaşları akut apandisit bağlamında makine öğrenmesi uygulamalarını ele alırken, Ma ve arkadaşları akut safra taşı pankreatiti şiddetinin erken öngörüsü için makine öğrenmesi modeli sunmuştur [64,142]. Koçer ve arkadaşları YOLOv5 ile abdominal BT görüntülerinde hastalık saptamayı ele almış, Park ve arkadaşları tekil ve seri BT görüntülerini karşılaştırarak akut abdominal hastalık sınıflandırmasında zamansal görüntü bilgisinin değerini tartışmıştır [151,188]."),
    ("p", "Bu grup, gelecekteki abdominal acil YZ sistemlerinin tek patoloji modellerinden çok çoklu patoloji, çoklu organ ve çoklu görev sistemlerine evrilebileceğini göstermektedir. Ancak bu geçiş için geniş etiketli veri setleri, klinik önceliklendirme mantığı, radyolog etkileşimi ve yanlış pozitif yönetimi birlikte ele alınmalıdır."),
]


def replace_section(doc: Document) -> None:
    paras = doc.paragraphs
    start = end = None
    for i, p in enumerate(paras):
        txt = p.text.strip()
        if txt.startswith("3.4. Patoloji"):
            start = i
        if start is not None and txt.startswith("3.5."):
            end = i
            break
    if start is None or end is None:
        raise RuntimeError(f"3.4 section bounds not found: {start}, {end}")

    # Insert new paragraphs before the old 3.4 heading, then remove old 3.4 block.
    anchor = paras[start]._p
    parent = anchor.getparent()
    insert_at = parent.index(anchor)
    for kind, text in section:
        p = doc.add_paragraph()
        if kind == "h2":
            p.style = "Heading 2"
        elif kind == "h3":
            p.style = "Heading 3"
        p.add_run(text)
        parent.remove(p._p)
        parent.insert(insert_at, p._p)
        insert_at += 1

    # Remove old paragraphs from start to just before original 3.5.
    for p in paras[start:end]:
        p._element.getparent().remove(p._element)


doc = Document(SRC)
replace_section(doc)
doc.save(OUT)
print(OUT.resolve())
