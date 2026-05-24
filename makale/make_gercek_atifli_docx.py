from docx import Document
from pathlib import Path


src = Path("BT_revize_analitik_detayli_TURKCE.docx")
out = Path("BT_revize_analitik_detayli_TURKCE_gercek_atifli.docx")

doc = Document(src)


def set_paragraph(idx: int, text: str) -> None:
    p = doc.paragraphs[idx]
    if not p.runs:
        p.add_run(text)
        return
    for run in p.runs:
        run.text = ""
    p.runs[0].text = text


original = {idx: doc.paragraphs[idx].text for idx in [18, 20, 22, 24, 30, 32, 38, 40, 128]}

set_paragraph(
    18,
    original[18].replace(
        "Bu yakınma, kendiliğinden sınırlanan benign durumlardan acil cerrahi veya girişimsel tedavi gerektiren yaşamı tehdit edici patolojilere kadar geniş bir klinik yelpazeyi temsil eder.",
        "Bu yakınma, kendiliğinden sınırlanan benign durumlardan acil cerrahi veya girişimsel tedavi gerektiren yaşamı tehdit edici patolojilere kadar geniş bir klinik yelpazeyi temsil eder [235].",
    ),
)

set_paragraph(
    20,
    original[20].replace(
        "Bilgisayarlı tomografi (BT), yüksek anatomik ayrıntı sağlaması ve kısa sürede tanısal bilgi sunması nedeniyle abdominal acil patolojilerde merkezi bir görüntüleme yöntemidir.",
        "Bilgisayarlı tomografi (BT), yüksek anatomik ayrıntı sağlaması ve kısa sürede tanısal bilgi sunması nedeniyle abdominal acil patolojilerde merkezi bir görüntüleme yöntemidir [235]-[237].",
    ),
)

set_paragraph(
    22,
    "Yapay zeka uygulamaları tıbbi görüntülemede saptama, sınıflandırma, bölütleme, lokalizasyon ve klinik karar destek süreçlerinde giderek daha fazla kullanılmaktadır [238]. "
    "Geleneksel makine öğrenmesi yaklaşımları çoğunlukla elle çıkarılmış radyomik özelliklere dayanırken, derin öğrenme yöntemleri ham görüntülerden hiyerarşik temsil öğrenme kapasitesiyle radyolojik örüntülerin uçtan uca modellenmesini sağlamaktadır [238], [239]. "
    "Son yıllarda evrişimli sinir ağları (CNN) ve U-Net türevlerinin yanında transformer tabanlı mimariler ve temel model yaklaşımları tıbbi görüntüleme alanında hızla benimsenmektedir [240]-[242].",
)

set_paragraph(24, original[24])

set_paragraph(
    30,
    "Bu çalışma JBI kapsam belirleme metodolojisi ile uyumlu olarak tasarlanmış ve raporlama PRISMA Extension for Scoping Reviews (PRISMA-ScR) kontrol listesine göre yapılmıştır [243], [244]. "
    "Çalışmaların nicel olarak birleştirilmesinden ziyade betimsel ve haritalayıcı olarak sentezlenmesi amaçlanmıştır; meta-analiz planlanmamıştır [243].",
)

set_paragraph(
    32,
    original[32]
    + " PCC çerçevesi kapsam belirleme incelemelerinde önerilen soru yapılandırma yaklaşımıyla uyumludur [243].",
)

set_paragraph(
    38,
    original[38]
    + " Bu akış, PRISMA-ScR raporlama ilkelerinde önerilen şeffaf seçim ve dışlama gerekçelendirme yaklaşımıyla uyumludur [244].",
)

set_paragraph(
    40,
    original[40]
    + " Veri çıkarım alanlarının önceden tanımlanması ve yinelemeli biçimde test edilmesi JBI kapsam belirleme rehberindeki önerilerle uyumludur [243].",
)

set_paragraph(
    128,
    "Aşağıdaki kaynakçanın ilk 234 referansı dahil edilen çalışmaların IEEE formatındaki kaynakçasıdır; 235-244 numaralı referanslar giriş ve yöntem bölümlerinde kullanılan genel klinik, teknik ve metodolojik kaynaklardır.",
)

new_refs = [
    '[235] Expert Panel on Gastrointestinal Imaging; C. D. Scheirey, K. J. Fowler, J. A. Therrien, D. H. Kim, W. B. Al-Refaie, M. A. Camacho, et al., "ACR Appropriateness Criteria Acute Nonlocalized Abdominal Pain," Journal of the American College of Radiology, vol. 15, no. 11S, pp. S217-S231, 2018. doi: 10.1016/j.jacr.2018.09.010',
    '[236] Expert Panel on Gastrointestinal Imaging, "ACR Appropriateness Criteria Right Lower Quadrant Pain-Suspected Appendicitis," Journal of the American College of Radiology, vol. 15, no. 11S, pp. S373-S387, 2018. doi: 10.1016/j.jacr.2018.09.033',
    '[237] G. Yarmish, M. Smith, M. Rosen, J. T. Baker, R. Blake, C. Cash, et al., "ACR Appropriateness Criteria Right Upper Quadrant Pain," Journal of the American College of Radiology, vol. 11, no. 3, pp. 316-322, 2014. doi: 10.1016/j.jacr.2013.11.017',
    '[238] G. Litjens, T. Kooi, B. E. Bejnordi, A. A. A. Setio, F. Ciompi, M. Ghafoorian, et al., "A survey on deep learning in medical image analysis," Medical Image Analysis, vol. 42, pp. 60-88, 2017. doi: 10.1016/j.media.2017.07.005',
    '[239] R. J. Gillies, P. E. Kinahan, and H. Hricak, "Radiomics: Images Are More than Pictures, They Are Data," Radiology, vol. 278, no. 2, pp. 563-577, 2016. doi: 10.1148/radiol.2015151169',
    '[240] O. Ronneberger, P. Fischer, and T. Brox, "U-Net: Convolutional Networks for Biomedical Image Segmentation," in Medical Image Computing and Computer-Assisted Intervention - MICCAI 2015, Lecture Notes in Computer Science, vol. 9351, pp. 234-241, 2015. doi: 10.1007/978-3-319-24574-4_28',
    '[241] A. Vaswani, N. Shazeer, N. Parmar, J. Uszkoreit, L. Jones, A. N. Gomez, L. Kaiser, and I. Polosukhin, "Attention Is All You Need," in Advances in Neural Information Processing Systems, vol. 30, 2017.',
    '[242] J. Ma, Y. He, F. Li, L. Han, C. You, and B. Wang, "Segment anything in medical images," Nature Communications, vol. 15, article 654, 2024. doi: 10.1038/s41467-024-44824-z',
    '[243] M. D. J. Peters, C. Marnie, A. C. Tricco, D. Pollock, Z. Munn, L. Alexander, P. McInerney, C. M. Godfrey, and H. Khalil, "Updated methodological guidance for the conduct of scoping reviews," JBI Evidence Synthesis, vol. 18, no. 10, pp. 2119-2126, 2020. doi: 10.11124/JBIES-20-00167',
    '[244] A. C. Tricco, E. Lillie, W. Zarin, K. K. O\'Brien, H. Colquhoun, D. Levac, et al., "PRISMA Extension for Scoping Reviews (PRISMA-ScR): Checklist and Explanation," Annals of Internal Medicine, vol. 169, no. 7, pp. 467-473, 2018. doi: 10.7326/M18-0850',
]

for ref in new_refs:
    doc.add_paragraph(ref)

doc.save(out)
print(out.resolve())
