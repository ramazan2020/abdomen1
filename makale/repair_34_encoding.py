from docx import Document
from pathlib import Path


path = Path("BT_revize_analitik_detayli_TURKCE.docx")
doc = Document(path)

start = None
end = None
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if start is None and text.startswith("3.4."):
        start = i
    elif start is not None and text.startswith("3.5."):
        end = i
        break

if start is None:
    raise RuntimeError("3.4 section not found")
if end is None:
    end = len(doc.paragraphs)

for p in doc.paragraphs[start:end]:
    try:
        fixed = p.text.encode("cp1252").decode("utf-8")
    except UnicodeError:
        fixed = p.text
    if fixed != p.text:
        p.text = fixed

doc.save(path)
print(path.resolve())
